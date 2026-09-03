from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("deliverables")
OUT_PATH = OUT_DIR / "IBPS_SIH_Internal_Hackathon_Detailed_Project_Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 34, 43)
MUTED = RGBColor(88, 96, 105)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ["top_margin", "right_margin", "bottom_margin", "left_margin"]:
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("IBPS SIH Internal Hackathon Report")
    set_run_font(run, size=9, color=MUTED)


def para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if len(str(value)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=9.5)
    set_table_geometry(table, widths)
    para(doc, "")
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    para(doc, "")


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("DETAILED PROJECT REPORT")
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("IBPS - Integrated Block Planning System")
    set_run_font(run, size=14, color=MUTED)

    for label, value in [
        ("Problem Statement", "PS 26027 - AI-Powered Automatic Block Planning to Maximize Asset Availability for Train Operations on Indian Railways"),
        ("Prepared For", "SIH Internal Hackathon Presentation Preparation"),
        ("Project Type", "Synthetic-data architectural prototype with CP-SAT optimization and React dashboard"),
        ("Primary Positioning", "AI-assisted decision support; final block approval remains with authorized railway personnel"),
        ("Generated On", date.today().strftime("%d %B %Y")),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=11, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.border_bottom = None
    r = p.add_run("This report is designed as the source document for creating the SIH internal hackathon PPT.")
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    doc.add_page_break()


def build():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_doc(doc)
    cover(doc)

    para(doc, "Executive Summary", "Heading 1")
    para(doc, "The Integrated Block Planning System (IBPS) is a decision-support prototype for coordinated railway maintenance block planning. It addresses the practical planning problem where Engineering, Signal & Telecommunication (S&T), and Traction Distribution (TRD) maintenance needs are raised through separate departmental processes, while train movements, block opportunities, freight forecasts, and control-office constraints exist in operational planning channels.")
    para(doc, "IBPS creates a unified planning layer. It accepts maintenance tasks, operational train movement data, goods traffic forecasts, and available block windows; scores the urgency of each task through an explainable priority engine; filters feasible task-to-block assignments through a candidate model; solves the planning problem using a CP-SAT constraint optimizer; compares the result with a fragmented baseline; and presents the plan, diagnostics, what-if impact, and KPIs through an API-backed React dashboard.")
    add_callout(doc, "One-line project pitch", "IBPS helps railway planners club compatible Engineering, S&T, and TRD work into safe, conflict-aware maintenance blocks so critical assets are restored faster while avoidable train disruption and wasted possession hours are reduced.")

    para(doc, "SIH Problem Alignment", "Heading 1")
    para(doc, "The project is aligned to Problem Statement 26027: AI-Powered Automatic Block Planning to Maximize Asset Availability for Train Operations on Indian Railways. The core business question implemented by the prototype is: What is the best feasible coordinated maintenance block plan for the chosen planning horizon?")
    bullet(doc, "Maintenance inputs cover task, department, asset, corridor, defect, severity, criticality, safety risk, overdue days, duration, crew, resources, precedence, and time bounds.")
    bullet(doc, "Operational inputs cover train movements, train priority, corridor, goods forecast, available block windows, and traffic density.")
    bullet(doc, "Planning outputs include scheduled tasks, unscheduled tasks, assigned block, start/end time, reason, conflict warnings, KPIs, and diagnostics.")
    bullet(doc, "Prototype boundaries are explicit: no live TMS, SMMS, TDMS, COA, or confidential railway database connection is claimed.")

    para(doc, "Problem Context", "Heading 1")
    para(doc, "Railway maintenance block planning is difficult because safety, operations, and maintenance all compete for the same scarce corridor time. Engineering may need track access, S&T may need signalling work, and TRD may need power-block conditions, but passenger and freight operations must continue with minimal disruption. When these requirements are planned in isolation, the system can suffer from repeated possessions, underused blocks, delayed critical maintenance, and weak visibility into why one activity was prioritized over another.")
    para(doc, "The project treats this as a coordination and optimization problem rather than only a dashboard problem. The system must combine task urgency, operational disruption, resource limits, safety permissions, precedence rules, block capacity, and clubbing benefits into one explainable plan.")
    add_table(doc, ["Pain Point", "Why It Matters", "IBPS Response"], [
        ("Departmental silos", "Engineering, S&T, and TRD plans may be generated independently.", "Unified domain model and cross-department clubbing objective."),
        ("Scarce block time", "Track possession opportunities are limited and operationally sensitive.", "Capacity-aware block assignment with block-hour penalty."),
        ("Safety criticality", "Critical defects cannot be treated like routine tasks.", "Priority engine, critical-task bonus, and unscheduled critical penalty."),
        ("Train conflicts", "Premium or high-priority train overlap may make a block infeasible.", "Candidate rejection for hard train conflicts and disruption penalties."),
        ("Planner trust", "Black-box scheduling is hard to accept in safety domains.", "Explainable score breakdowns, rejection reasons, diagnostics, and human approval notice."),
    ], [1900, 3400, 4060])

    para(doc, "Solution Overview", "Heading 1")
    para(doc, "IBPS is organized around a repeatable planning pipeline:")
    number(doc, "Ingest maintenance and operational data through adapters.")
    number(doc, "Normalize records into typed Pydantic domain models.")
    number(doc, "Compute explainable priority scores for each maintenance task.")
    number(doc, "Evaluate every task-block pair for corridor, time, duration, department, safety, crew, and train-conflict feasibility.")
    number(doc, "Build and solve a CP-SAT model with hard constraints and soft objective terms.")
    number(doc, "Compare the optimized plan with a siloed greedy baseline.")
    number(doc, "Expose results through API endpoints and a React dashboard.")
    number(doc, "Support emergency what-if replanning and show the exact schedule diff.")

    para(doc, "Stakeholders and Users", "Heading 1")
    add_table(doc, ["Stakeholder", "Need", "How IBPS Supports It"], [
        ("Block planner / controller", "A feasible possession plan with reasons.", "Optimized schedule, conflict warnings, diagnostic reasons, and human override framing."),
        ("Engineering", "Critical track and civil defects cleared within safe windows.", "Severity, safety, overdue, deadline, crew, and resource constraints."),
        ("S&T", "Signal and telecom work coordinated with track and TRD conditions.", "Department permissions, resource capacity, and shared block visibility."),
        ("TRD", "Power block constraints and OHE safety conditions respected.", "Safety constraints such as POWER_BLOCK_AVAILABLE and permitted departments."),
        ("Operations", "Train movement and freight impact considered.", "Hard train-conflict pruning and disruption cost terms."),
        ("Hackathon evaluator", "Evidence that the prototype is real and explainable.", "Live APIs, tests, diagnostics, synthetic-data badge, and measurable KPI comparison."),
    ], [2050, 2850, 4460])

    para(doc, "System Architecture", "Heading 1")
    para(doc, "The architecture follows a clean dependency direction: UI to API to application/domain services to optimizer, scoring, metrics, and data. The optimizer remains independent of FastAPI and React so it can be tested as a pure planning engine.")
    add_table(doc, ["Layer", "Repository Area", "Responsibility"], [
        ("Data sources/adapters", "backend/app/adapters, backend/app/data", "Synthetic TMS, SMMS, TDMS, and COA adapter interfaces plus deterministic demo/full datasets."),
        ("Domain model", "backend/app/domain", "Typed models for tasks, trains, goods forecasts, block windows, candidates, plans, metrics, and replan diffs."),
        ("Priority scoring", "backend/app/scoring", "Explainable rule-based urgency score and priority band classification."),
        ("Optimization", "backend/app/optimization", "Candidate feasibility model, CP-SAT solver, baseline scheduler, objective terms, and diagnostics."),
        ("Metrics", "backend/app/metrics", "KPIs and baseline-vs-optimized comparison, including simulated availability proxy."),
        ("Services/API", "backend/app/services, backend/app/api", "State management, route handlers, dataset switching, what-if replanning, health, diagnostics, plans, blocks, and tasks."),
        ("Frontend", "frontend/src", "React dashboard with Overview, Tasks, Blocks, Planning, What-if, and Diagnostics pages."),
    ], [1900, 2600, 4860])

    para(doc, "Data Model", "Heading 1")
    para(doc, "The central model is a maintenance task with department, asset, corridor, defect type, severity, criticality, safety risk, overdue days, duration, crew requirement, resource requirements, dependencies, incompatible tasks, earliest start, deadline, status, traffic criticality, and computed priority fields. Supporting models represent train movements, freight forecasts, block windows, candidate evaluations, scheduled tasks, plan metrics, and replan diffs.")
    add_table(doc, ["Entity", "Key Fields", "Purpose"], [
        ("MaintenanceTask", "department, asset, corridor, defect, severity, risk, duration, crew, precedence, deadline", "Represents work demand from Engineering, S&T, or TRD."),
        ("BlockWindow", "corridor, start/end, available capacity, resource capacity, safety constraints, permitted departments", "Represents a possession opportunity where work can be scheduled."),
        ("TrainMovement", "corridor, train type, time interval, operational priority, disruption penalty", "Represents operating traffic that may conflict with blocks."),
        ("GoodsForecast", "corridor, time window, expected goods trains, probability, traffic density", "Models freight traffic pressure by time window."),
        ("BlockPlan", "scheduled tasks, unscheduled tasks, blocks used, metrics, warnings, solver status", "Stores the output of baseline or optimized planning."),
        ("SolverDiagnosticReport", "candidate counts, rejection tallies, objective terms, unscheduled reasons", "Gives planners and judges explainability for solver behavior."),
    ], [1950, 4300, 3110])

    para(doc, "Synthetic Dataset Strategy", "Heading 1")
    para(doc, "The prototype uses a deterministic demo fixture and a scalable synthetic dataset. The demo fixture is crafted for the internal hackathon story: it contains representative Engineering, S&T, and TRD tasks, multiple corridors, premium train conflicts, freight forecasts, block windows, critical tasks, precedence, safety constraints, and emergency replanning. The full dataset provides a scaled synthetic mode with about 200 tasks across 12 trunk corridors, 48 blocks, and 48 trains.")
    bullet(doc, "Demo fixture: 21 deterministic maintenance tasks designed for evaluator-visible scenarios.")
    bullet(doc, "Full dataset: scaled synthetic mode for demonstrating architectural scalability.")
    bullet(doc, "All data is labeled as synthetic/demo data only; no confidential railway data is used.")

    para(doc, "Priority Engine", "Heading 1")
    para(doc, "IBPS uses an explainable rule-based priority engine rather than black-box machine learning. This choice is intentional because safety-critical planning needs traceability. Every task receives a 0 to 100 priority score and a band: CRITICAL, HIGH, MEDIUM, or ROUTINE.")
    add_table(doc, ["Factor", "Default Weight", "Interpretation"], [
        ("Defect severity", "0.25", "Maps severity to a normalized value: Critical 100, Major 75, Minor 45, Routine 20."),
        ("Asset criticality", "0.25", "Captures the operational importance of the asset."),
        ("Safety risk", "0.25", "Represents hazard severity and safety urgency."),
        ("Overdue days", "0.15", "Uses min(100, overdue days x 12.5), so delay quickly raises urgency."),
        ("Traffic criticality", "0.10", "Gives extra weight to work affecting important corridors."),
    ], [2600, 1800, 4960])
    para(doc, "Priority Score = 0.25*Severity + 0.25*Criticality + 0.25*SafetyRisk + 0.15*OverdueScore + 0.10*TrafficCriticality")
    para(doc, "Priority bands are CRITICAL for scores >= 80, HIGH for >= 60, MEDIUM for >= 40, and ROUTINE below 40. The frontend can display the score breakdown, helping evaluators see why a critical overdue track defect is ranked above routine work.")

    para(doc, "Candidate Feasibility Model", "Heading 1")
    para(doc, "The candidate model evaluates whether a task can legally and practically fit into a block before the solver sees the decision variable. This reduces solver size and makes rejection reasons explainable.")
    bullet(doc, "Corridor must match between task and block.")
    bullet(doc, "Department must be permitted in the block.")
    bullet(doc, "Task duration must fit inside the block window.")
    bullet(doc, "Block timing must respect the task earliest start and deadline.")
    bullet(doc, "Crew required by the task must fit within block resource capacity.")
    bullet(doc, "TRD/OHE work must have required safety constraints such as POWER_BLOCK_AVAILABLE.")
    bullet(doc, "Priority-1 train overlap, including premium train movement, is treated as a hard train conflict and rejected.")

    para(doc, "CP-SAT Optimization Model", "Heading 1")
    para(doc, "The optimizer creates binary decision variables x[i,j], where x[i,j] = 1 means maintenance task i is assigned to block j. A task may remain unscheduled if no feasible or high-value allocation exists. Auxiliary variables track whether a task is scheduled, whether a block is used, and whether departments are active in a block.")
    add_table(doc, ["Constraint", "Hard Rule Implemented"], [
        ("Task uniqueness", "Each task is assigned to at most one block."),
        ("Block capacity", "Number of tasks in a block cannot exceed available task slots."),
        ("Crew/resource capacity", "Sum of required crew in a block cannot exceed resource capacity."),
        ("Time bounds", "Assignment must respect earliest start and deadline."),
        ("Precedence", "Dependent task cannot be scheduled unless predecessor is scheduled first."),
        ("Safety incompatibility", "Incompatible tasks cannot share the same block."),
        ("Corridor and department", "Only corridor-compatible and permitted-department candidates are created."),
        ("Premium train conflict", "Hard conflict candidates are pruned before optimization."),
    ], [2550, 6810])
    para(doc, "The soft objective maximizes completed priority value, gives extra reward for critical tasks, penalizes unscheduled critical work, rewards cross-department clubbing, penalizes block hours, and penalizes train/freight disruption. This lets the system prefer high-value safe coordination rather than simply filling every available slot.")
    add_table(doc, ["Objective Term", "Planning Behavior Encouraged"], [
        ("Priority completion", "Schedule high-value work first."),
        ("Critical completion bonus", "Protect safety-critical defect clearance."),
        ("Unscheduled critical penalty", "Avoid leaving urgent safety work out of the plan."),
        ("Clubbing bonus", "Encourage compatible Engineering, S&T, and TRD work in the same possession."),
        ("Block-hour penalty", "Avoid unnecessary corridor closures."),
        ("Train disruption penalty", "Prefer lower-disruption windows."),
        ("Goods traffic penalty", "Avoid freight-heavy periods where possible."),
        ("What-if soft pinning", "During replanning, preserve stable assignments where useful."),
    ], [2550, 6810])

    para(doc, "Baseline Comparison", "Heading 1")
    para(doc, "The baseline scheduler simulates a fragmented planning practice: departments are processed sequentially, each department selects feasible blocks greedily, and cross-department optimization is not considered. This gives a meaningful reference point for the optimized plan.")
    add_table(doc, ["Metric", "Baseline", "Optimized IBPS", "Reported Improvement"], [
        ("Critical safety defects cleared", "4 tasks", "4 tasks", "No change; both clear critical tasks in demo."),
        ("Priority score fulfilled", "76.6%", "94.1%", "+22.9% relative improvement."),
        ("Cross-department coordinated blocks", "3 blocks", "5 blocks", "+66.7% more coordinated blocks."),
        ("Average block utilization", "76.4%", "94.4%", "+23.6% relative improvement."),
        ("Simulated asset availability", "95.3%", "96.2%", "+0.9 percentage points in the demo proxy."),
    ], [2600, 1600, 1800, 3360])
    para(doc, "These figures are simulation outputs from the prototype dataset, not claims from a production railway deployment. The report and UI should keep the synthetic-data badge visible.")

    para(doc, "Metrics and KPIs", "Heading 1")
    para(doc, "IBPS computes KPIs from actual schedule output, not manually entered claims. Plan metrics include total tasks, scheduled tasks, unscheduled tasks, total critical tasks, critical tasks completed, total and completed priority score, priority fulfillment percentage, blocks used, total block hours, train conflict count, total train disruption penalty, goods traffic penalty, average block utilization, multi-department clubbed block count, and simulated asset availability.")
    para(doc, "Simulated Asset Availability = 80.0% + 12.0%*(Critical Completed / Total Critical) + 5.0%*(Completed Priority / Total Priority) - 2.5%*(Unscheduled Critical) - 0.5%*(Train Conflicts). The value is clamped between 45.0% and 99.9%. This is intentionally labeled as a proxy, not an official railway asset availability formula.")

    para(doc, "What-If Emergency Replanning", "Heading 1")
    para(doc, "The what-if module demonstrates how IBPS reacts when a new safety-critical defect appears after a plan already exists. The demo emergency task is injected, scored, and included in a new optimization run. The output does not merely return a new schedule; it returns a diff showing tasks added, moved, displaced, and unchanged.")
    bullet(doc, "Example: a sudden transverse rail crack is inserted as a critical Engineering task.")
    bullet(doc, "The optimizer accommodates the emergency task if a feasible high-value block exists.")
    bullet(doc, "Lower-priority work may be displaced when capacity must be reserved for the emergency task.")
    bullet(doc, "The UI shows the before/after KPI impact and exact plan changes, preserving planner trust.")

    para(doc, "Backend API Surface", "Heading 1")
    add_table(doc, ["Area", "Endpoint Examples", "Purpose"], [
        ("System", "GET /api/health, GET /api/dashboard", "Health, version, synthetic data mode, overview metrics."),
        ("Tasks", "GET /api/tasks, GET /api/tasks/{task_id}", "Task list, filtering, detail, score breakdown, candidate reasons."),
        ("Blocks", "GET /api/blocks, GET /api/blocks/{block_id}", "Block windows, utilization, assigned departments, conflict context."),
        ("Plans", "POST /api/plans/baseline, POST /api/plans/optimize, GET /api/plans/comparison", "Generate and compare baseline and optimized plans."),
        ("What-if", "POST /api/plans/what-if", "Emergency defect injection and replan diff."),
        ("Diagnostics", "GET /api/plans/diagnostics", "Candidate counts, rejection tally, objective terms, unscheduled reasons."),
        ("Datasets", "POST /api/datasets/switch", "Switch between demo fixture and full synthetic dataset."),
    ], [1600, 3300, 4460])

    para(doc, "Frontend Experience", "Heading 1")
    para(doc, "The frontend is a Vite React TypeScript application. It uses route-based pages, typed API clients, reusable UI components, Lucide icons, and Recharts charts. The design is built around an operational dashboard rather than a marketing landing page.")
    add_table(doc, ["Screen", "What It Shows", "How To Use In PPT Demo"], [
        ("Overview", "Asset availability proxy, plan comparison, priority tasks, upcoming blocks, operational insights.", "Open with the main before/after value proposition."),
        ("Tasks", "Maintenance workbank with filters, priority badges, status, detail and feasibility explanation.", "Show explainability: why a task is urgent and why blocks are accepted/rejected."),
        ("Blocks", "Block windows, capacity utilization, assigned departments, safety constraints.", "Show clubbing of Engineering, S&T, and TRD work."),
        ("Planning", "Generate baseline and optimized plans, compare solver results.", "Demonstrate the optimizer solving live rather than using static screenshots."),
        ("What-if", "Emergency defect form, replan result, added/moved/displaced tasks.", "Use as the climax of the demo story."),
        ("Diagnostics", "Candidate-pair counts, rejection reasons, objective breakdown, unscheduled explanations.", "Use when judges ask how the AI decision is trustworthy."),
    ], [1600, 3600, 4160])

    para(doc, "Testing and Validation", "Heading 1")
    para(doc, "The repository contains unit, API, and integration tests covering the core planning logic and demo scenarios. This is important for the PPT because it proves the project is not only a UI mockup.")
    bullet(doc, "Priority engine tests validate score calculation and routine/critical banding.")
    bullet(doc, "Candidate model tests validate corridor mismatch and train conflict rejection.")
    bullet(doc, "Optimizer constraint tests validate at-most-once assignment, corridor isolation, deadline enforcement, resource capacity, precedence, incompatible tasks, missing safety constraint rejection, and train conflict prohibition.")
    bullet(doc, "Objective tests validate cross-department clubbing incentive and sensitivity across objective weight profiles.")
    bullet(doc, "Metrics tests validate live KPI computation.")
    bullet(doc, "Baseline tests validate the greedy fragmented scheduler.")
    bullet(doc, "What-if tests validate emergency insertion and displacement behavior.")
    bullet(doc, "API tests validate health, dashboard, tasks, blocks, plans, comparison, diagnostics, what-if, deterministic optimization runs, and dataset switching.")
    bullet(doc, "End-to-end demo story test validates the full demo fixture flow.")

    para(doc, "Demo Storyboard", "Heading 1")
    para(doc, "A strong SIH demo should feel like a realistic control-room planning story, not a feature checklist.")
    number(doc, "Start on the Overview screen: show synthetic-data badge, asset availability proxy, priority completion, coordinated blocks, and critical defects cleared.")
    number(doc, "Open Tasks: filter or sort to show critical/high tasks and explain the priority formula.")
    number(doc, "Open a Task detail: show feasible and rejected block candidates with reasons such as train conflict, corridor mismatch, deadline, or safety constraint.")
    number(doc, "Open Blocks: show how multiple departments are clubbed into one coordinated block where safety and capacity allow.")
    number(doc, "Open Planning: generate or compare baseline versus optimized plans and explain CP-SAT in simple terms.")
    number(doc, "Open Diagnostics: show candidate pairs, rejection tally, objective contributions, and unscheduled task reasons.")
    number(doc, "Open What-if: inject an emergency defect and show added/displaced/unchanged tasks in the replan diff.")
    number(doc, "Close with human-in-the-loop message: the system recommends; authorized railway personnel approve.")

    para(doc, "Innovation Value", "Heading 1")
    bullet(doc, "Integrated block planning across Engineering, S&T, and TRD instead of siloed departmental scheduling.")
    bullet(doc, "Explainable priority scoring designed for safety-critical planning transparency.")
    bullet(doc, "True constraint optimization using CP-SAT with hard constraints and multi-objective soft terms.")
    bullet(doc, "Cross-department clubbing incentive to reduce repeated corridor possession.")
    bullet(doc, "Diagnostic transparency: rejected candidates, unscheduled reasons, objective contribution breakdown, and KPI deltas.")
    bullet(doc, "Emergency what-if replanning that highlights schedule churn and displacement impact.")
    bullet(doc, "Adapter-based architecture ready for future integration with TMS, SMMS, TDMS, COA, timetable, and freight forecasts.")

    para(doc, "Feasibility and Scalability", "Heading 1")
    para(doc, "The current system is feasible as a prototype because the optimization problem is modeled with explicit binary variables and candidate pruning. Candidate pruning avoids creating impossible decision variables, keeping the solver problem smaller. The state service supports both a small deterministic fixture for demos and a larger synthetic dataset for scale demonstration. FastAPI separates the optimization engine from UI concerns, and React consumes typed API contracts.")
    para(doc, "For production, the same architecture would need authenticated data adapters, stronger security, audit logging, operational validation by railway experts, integration testing with real data feeds, approval workflows, and monitored model/solver performance. The current project intentionally stops at architectural prototype and decision-support demonstration.")

    para(doc, "Risks, Limitations, and Honesty Notes", "Heading 1")
    add_table(doc, ["Limitation", "Current Position", "Future Work"], [
        ("Synthetic data", "All records are generated for demonstration.", "Replace synthetic adapters with authorized live or anonymized data feeds."),
        ("Availability proxy", "Transparent simulated KPI, not official IR metric.", "Calibrate with accepted railway asset availability definitions."),
        ("Operational approval", "IBPS recommends plans only.", "Add approval workflow, audit trail, override reason capture, and role-based permissions."),
        ("Monthly planning", "Architecture supports horizon concept; weekly detailed planning is the strongest current demo.", "Add aggregation strategies and larger-horizon performance tuning."),
        ("Safety rule depth", "Prototype encodes representative constraints.", "Validate a complete rulebook with domain experts."),
        ("Solver performance", "Current datasets are suitable for demo scale.", "Tune time limits, decomposition, and rolling-horizon solves for production scale."),
    ], [2200, 3400, 3760])

    para(doc, "Roadmap", "Heading 1")
    add_table(doc, ["Phase", "Goal", "Major Additions"], [
        ("Phase 1 - Prototype", "Demonstrate integrated block planning on synthetic data.", "Priority engine, CP-SAT optimizer, baseline comparison, dashboard, what-if, diagnostics."),
        ("Phase 2 - Pilot data readiness", "Connect controlled sample data and validate with experts.", "Adapter hardening, data quality checks, calibration of scoring weights, richer safety rules."),
        ("Phase 3 - Planner workflow", "Support real planner review and controlled overrides.", "Role-based access, approval workflow, audit log, comments, scenario versioning."),
        ("Phase 4 - Operational integration", "Move toward assisted live planning.", "Timetable/COA feeds, notifications, rescheduling hooks, production monitoring."),
        ("Phase 5 - Advanced intelligence", "Improve robustness and predictive planning.", "Delay prediction, reliability modeling, learned disruption cost calibration, rolling-horizon optimization."),
    ], [1900, 2900, 4560])

    para(doc, "PPT Conversion Guide", "Heading 1")
    para(doc, "Suggested 12-slide structure:")
    add_table(doc, ["Slide", "Title", "Message"], [
        ("1", "IBPS: Integrated Block Planning System", "AI-assisted coordinated block planning for Indian Railways maintenance."),
        ("2", "Problem: Planning Across Silos", "Engineering, S&T, TRD, and operations compete for scarce block windows."),
        ("3", "Our Solution", "A unified planning layer that scores, filters, optimizes, compares, and explains."),
        ("4", "System Architecture", "Adapters, unified domain model, priority engine, candidate model, CP-SAT optimizer, metrics, API, React dashboard."),
        ("5", "Data and Boundary Honesty", "Synthetic data prototype with clear future adapter interfaces and human approval."),
        ("6", "Priority Scoring", "Explainable urgency score using severity, criticality, safety risk, overdue days, and traffic criticality."),
        ("7", "Optimization Model", "Hard constraints protect feasibility; objective rewards high-priority completion and cross-department clubbing."),
        ("8", "Dashboard Walkthrough", "Overview, task explainability, block utilization, plan comparison, diagnostics."),
        ("9", "Demo: Baseline vs IBPS", "IBPS improves priority fulfillment, coordination, utilization, and simulated availability in the demo fixture."),
        ("10", "Demo: Emergency What-if", "Inject a critical defect and show added/displaced/unchanged tasks."),
        ("11", "Validation", "Unit, API, integration, objective, constraint, metrics, and end-to-end tests."),
        ("12", "Impact and Roadmap", "Reduced planning fragmentation, safer prioritization, better utilization, and path to pilot integration."),
    ], [900, 2950, 5510])

    para(doc, "Suggested Judge Q&A", "Heading 1")
    add_table(doc, ["Question", "Strong Answer"], [
        ("Is this using real railway data?", "No. It is a synthetic-data architectural prototype. The architecture is designed for future authorized adapters to TMS, SMMS, TDMS, COA, timetable, and freight forecasts."),
        ("Where is AI used?", "In this prototype, AI is represented as algorithmic decision support: explainable priority scoring, constraint optimization, what-if replanning, and diagnostic recommendations. We avoid black-box ML on the safety-critical path."),
        ("How do you ensure safety?", "Safety-critical conditions are modeled as hard constraints where appropriate: safety permissions, incompatibilities, train conflicts, capacity, crew limits, time bounds, and precedence. Final approval stays with authorized personnel."),
        ("Why CP-SAT?", "The problem is a constrained assignment and scheduling problem with binary decisions, hard rules, and weighted objectives. CP-SAT is well suited for feasibility plus optimization."),
        ("What is the main impact?", "Better use of block windows through coordinated cross-department clubbing, higher priority-score fulfillment, transparent decision reasons, and emergency replanning visibility."),
        ("What remains before production?", "Validated data integration, real operational rulebooks, security, audit trail, approval workflow, expert calibration, and production-scale testing."),
    ], [2500, 6860])

    para(doc, "Recommended Presentation Line", "Heading 1")
    para(doc, "IBPS does not replace railway planners. It gives planners a mathematically optimized, explainable, and auditable recommendation so that scarce maintenance blocks can be used for the most urgent and compatible work while protecting train operations and safety approvals.")

    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    build()
